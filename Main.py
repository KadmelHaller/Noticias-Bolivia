import streamlit as st
import google.generativeai as genai
import requests
from bs4 import BeautifulSoup
from datetime import datetime
import pytz
from urllib.parse import urlparse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import time

# --- CONFIGURACIÓN ---
st.set_page_config(page_title="Monitor Estratégico Bolivia", page_icon="🇧🇴", layout="wide")
zona_horaria = pytz.timezone('America/La_Paz')
fecha_hoy = datetime.now(zona_horaria).strftime('%d/%m/%Y')

KEYWORDS = ["impuesto", "sin", "tribut", "factur", "fiscal", "recauda", "aduana", "gobierno", "arce", "ministro", "estado", "economia", "dolar", "banco", "crisis"]

if 'datos_medios' not in st.session_state: st.session_state.datos_medios = ""
if 'datos_rrss' not in st.session_state: st.session_state.datos_rrss = ""
if 'reporte_final' not in st.session_state: st.session_state.reporte_final = ""

st.title(f"🔍 MONITOR ESTRATÉGICO TOTAL: {fecha_hoy}")
api_key = st.sidebar.text_input("Ingresa tu Gemini API Key:", type="password")

# --- FUNCIONES DE SOPORTE ---

def obtener_info_envio():
    ahora = datetime.now(zona_horaria)
    hm = ahora.hour * 100 + ahora.minute
    if 830 <= hm <= 930: return "1er Envío"
    elif 1330 <= hm <= 1430: return "2do Envío"
    elif 1530 <= hm <= 1630: return "3er Envío"
    return "Envío Extraordinario"

def scraping_seguro(fuentes):
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) Chrome/122.0.0.0 Safari/537.36', 'Referer': 'https://www.google.com/'}
    acumulado = ""
    pb = st.progress(0)
    for i, f in enumerate(fuentes):
        st.write(f"📡 Buscando en: {f['n']}...")
        try:
            r = requests.get(f['u'], headers=headers, timeout=15)
            if r.status_code == 200:
                soup = BeautifulSoup(r.text, 'html.parser')
                elementos = soup.find_all(['a', 'h1', 'h2', 'h3'])
                encontrados_en_fuente = 0
                for el in elementos:
                    tit = el.get_text().strip()
                    url = el.get('href', '')
                    if len(tit) > 25 and any(k in (tit + url).lower() for k in KEYWORDS):
                        acumulado += f"FUENTE: {f['n']} | CIUDAD: {f['r']} | TITULAR: {tit} | URL: {url}\n\n"
                        encontrados_en_fuente += 1
                if encontrados_en_fuente == 0:
                    st.write(f"⚠️ No se hallaron coincidencias en {f['n']}.")
            else:
                st.write(f"❌ Error {r.status_code} en {f['n']}")
        except Exception as e:
            st.write(f"❌ Error de conexión en {f['n']}")
        pb.progress((i + 1) / len(fuentes))
        time.sleep(1) # Pausa para evitar bloqueos
    return acumulado

def procesar_ia_total(datos_prensa, datos_rrss):
    try:
        modelos = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        modelo_final = next((m for m in modelos if 'flash' in m), modelos[0])
        model = genai.GenerativeModel(modelo_final)
        
        prompt = f"""
        FECHA: {fecha_hoy}.
        TAREA: Generar un reporte unificado de prensa, TV y redes sociales.
        PRIORIDAD: 1. IMPUESTOS, 2. GOBIERNO, 3. ECONOMÍA.

        INSTRUCCIONES DE FIDELIDAD:
        - TITULARES: Deben ser EXACTAMENTE los que aparecen en la fuente, en MAYÚSCULAS.
        - PERSONAS: Nombres y cargos exactos (ej. "Mario Cazón, Presidente del SIN").
        - RRSS: Si hay datos de redes sociales, lístalos dentro de la ciudad que corresponda.

        FORMATO POR NOTICIA:
        *TITULAR EXACTO*
        NOMBRE DEL MEDIO O RED SOCIAL
        Resumen de 5 líneas sin etiquetas.
        Link directo.

        ESTRUCTURA: Divide solo en 1. COCHABAMBA y 2. SANTA CRUZ.
        REGLA: Usa UN solo asterisco (*) para el titular. El medio en la línea de abajo.
        """
        datos_combinados = f"--- PRENSA Y TV ---\n{datos_prensa}\n\n--- REDES SOCIALES ---\n{datos_rrss}"
        res = model.generate_content(prompt + "\n\nDATOS:\n" + datos_combinados)
        return res.text
    except Exception as e:
        return f"Error en IA: {str(e)}"

def generar_word_bolivia(texto, ciudad_tag):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    doc.add_paragraph("N°")
    doc.add_paragraph(fecha_hoy)
    doc.add_paragraph(obtener_info_envio())
    doc.add_paragraph("")

    target = "COCHABAMBA" if ciudad_tag == "CBBA" else "SANTA CRUZ"
    lineas = texto.split('\n')
    capturando = False
    
    for linea in lineas:
        if target in linea.upper() and ("1." in linea or "2." in linea):
            capturando = True
            continue
        if capturando and any(x in linea.upper() for x in ["1. COCHABAMBA", "2. SANTA CRUZ"]) and target not in linea.upper():
            capturando = False
            
        if capturando and linea.strip():
            if linea.startswith('*') and linea.endswith('*'):
                p = doc.add_paragraph()
                p.add_run(linea.replace('*', '').upper()).bold = True
            else:
                doc.add_paragraph(linea)
    
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- FLUJO DE TRABAJO ---

if api_key:
    genai.configure(api_key=api_key)
    
    # PASO 1: MEDIOS
    if st.button('🚀 PASO 1: REPORTE DE MEDIOS (PRENSA/TV/DIGITAL)'):
        fuentes_m = [
            {"n": "OPINIÓN", "u": "https://www.opinion.com.bo/", "r": "Cochabamba"},
            {"n": "LOS TIEMPOS", "u": "https://www.lostiempos.com/", "r": "Cochabamba"},
            {"n": "INNOTICIAS", "u": "https://innoticiasbo.com/", "r": "Cochabamba"},
            {"n": "ENFOQUE NEWS", "u": "https://enfoquenews.com.bo/", "r": "Cochabamba"},
            {"n": "EL DEBER", "u": "https://eldeber.com.bo/", "r": "Santa Cruz"},
            {"n": "EL MUNDO", "u": "https://elmundo.com.bo/", "r": "Santa Cruz"},
            {"n": "LA VOZ DIGITAL", "u": "https://lavoz.digital/", "r": "Santa Cruz"},
            {"n": "VISIÓN 360", "u": "https://www.vision360.bo/", "r": "Nacional"},
            {"n": "UNITEL", "u": "https://unitel.bo/", "r": "Nacional"},
            {"n": "RED UNO", "u": "https://www.reduno.com.bo/", "r": "Nacional"}
        ]
        st.session_state.datos_medios = scraping_seguro(fuentes_m)
        st.session_state.reporte_final = procesar_ia_total(st.session_state.datos_medios, "")
        st.success("Reporte de medios generado.")

    # PASO 2: REDES SOCIALES
    if st.session_state.datos_medios:
        if st.button('📱 PASO 2: AGREGAR REDES SOCIALES'):
            rrss_q = [
                {"n": "RRSS CBBA", "u": "https://www.google.com/search?q=impuestos+Cochabamba+site:facebook.com+OR+site:x.com&tbs=qdr:d", "r": "Cochabamba"},
                {"n": "RRSS SCZ", "u": "https://www.google.com/search?q=impuestos+Santa+Cruz+site:facebook.com+OR+site:x.com&tbs=qdr:d", "r": "Santa Cruz"}
            ]
            st.session_state.datos_rrss = scraping_seguro(rrss_q)
            if not st.session_state.datos_rrss:
                st.info("No se hallaron nuevas publicaciones en RRSS en las últimas 24h.")
            
            # Actualizamos el reporte integrando RRSS
            st.session_state.reporte_final = procesar_ia_total(st.session_state.datos_medios, st.session_state.datos_rrss)
            st.success("Redes sociales integradas al reporte.")

    # PASO 3: VISUALIZACIÓN Y EXPORTACIÓN
    if st.session_state.reporte_final:
        st.markdown("### 📄 VISTA PREVIA DEL REPORTE UNIFICADO")
        st.markdown(f'<div style="background:white;color:black;padding:25px;border:1px solid #ccc;font-family:serif;">{st.session_state.reporte_final.replace("\n", "<br>")}</div>', unsafe_allow_html=True)
        
        st.markdown("---")
        st.subheader("💾 PASO 3: EXPORTAR DOCUMENTOS")
        c1, c2 = st.columns(2)
        with c1:
            w_cbba = generar_word_bolivia(st.session_state.reporte_final, "CBBA")
            st.download_button("📄 Word Cochabamba (CBBA)", data=w_cbba, file_name=f"Reporte_{fecha_hoy.replace('/','-')}_CBBA.docx")
        with c2:
            w_scz = generar_word_bolivia(st.session_state.reporte_final, "STACRUZ")
            st.download_button("📄 Word Santa Cruz (STACRUZ)", data=w_scz, file_name=f"Reporte_{fecha_hoy.replace('/','-')}_STACRUZ.docx")
else:
    st.info("Ingresa tu API Key para comenzar.")
